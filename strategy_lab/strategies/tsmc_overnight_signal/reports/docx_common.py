#!/usr/bin/env python3
"""
TSMC 隔夜訊號 Word 報告的共用零件（樣式、表格、圖片、CSV 讀取）。

**單獨成檔是為了解掉 `generate_docx` ↔ `docx_append` 的循環 import**
（健檢 F-006）：原本兩檔互相 import，靠 import 順序與函式內延遲 import 僥倖
不炸。更實際的代價是它讓 `scripts/check_layer_deps.py` 的結束碼永遠是 1，
於是這支專門為了守住分層而寫的腳本沒辦法接進 CI 或 pre-commit 當閘門——
**唯一的違規把整個閘門鎖住了**。

拆法：兩檔共用的零件（`_OUTPUT` 路徑、`sanitize_for_word()`、樣式設定、
表格／圖片產生器、`_format_metrics_rows()`、`_method_usage_rows()`）搬到本檔，
`generate_docx` 只留 `build_report()` 與 CLI，`docx_append` 只留敘事內容，
兩者都單向 import 本檔。
"""

from __future__ import annotations

import datetime as dt
import re
from pathlib import Path
from typing import Dict, List, Literal, Optional

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as e:
    raise SystemExit("請先安裝 python-docx：pip install python-docx\n" + str(e)) from e

# reports/ 與 output/ 是 tsmc_overnight_signal/ 底下的同層姊妹資料夾
_STRATEGY_DIR = Path(__file__).resolve().parent.parent
_OUTPUT = _STRATEGY_DIR / "output"
_DOCX_ZH = Path(__file__).resolve().parent / "TSMC_OvernightSignal_Quant_Report.docx"
_DOCX_EN = Path(__file__).resolve().parent / "TSMC_OvernightSignal_Quant_Report_EN.docx"


Lang = Literal["zh", "en"]


def sanitize_for_word(text: str) -> str:
    """
    過濾易在 Word 造成顯示異常或與 LaTeX 混淆的字元。
    - 移除 $（避免被誤認為數學模式）
    - 替換全形／特殊連字與異常空白
    - 將 Unicode 減號統一為 ASCII 連字號（表格與數字可讀性）
    """
    if not text:
        return text
    s = text.replace("$", "")
    s = s.replace("\u00a0", " ")  # nbsp
    s = s.replace("\u2028", " ")
    s = s.replace("\u2029", " ")
    s = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", s)  # zero-width
    s = s.replace("\u2212", "-").replace("\u2013", "-").replace("\u2014", "-")
    s = s.replace("§", "")
    return s.strip()


# ---------------------------------------------------------------------------
# 標題與用語（統一使用隔夜訊號／跨市場領先）
# ---------------------------------------------------------------------------
TITLES_ZH = {
    "title": "台積電跨市場定價與隔夜訊號研究",
    "subtitle": "ADR、費城半導體與匯率之領先資訊：預測、執行假設與回測穩健性",
    "tagline": "量化交易期末專題報告 · AlphaEdge",
}

TITLES_EN = {
    "title": "TSMC Cross-Market Pricing and Overnight Information",
    "subtitle": "Lead Signals from ADR, SOX, and FX: Forecasting, Execution, and Backtest Robustness",
    "tagline": "Quantitative Trading Course Final Report · AlphaEdge",
}


def _ensure_style_eastasia(style, east: str = "Microsoft JhengHei") -> None:
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), east)


def _style_set_latin_font(
    style, ascii_font: str, east_asia: Optional[str] = None
) -> None:
    """Set Word rFonts so body text renders with the intended Latin font (and matching eastAsia for Word UI)."""
    east = east_asia or ascii_font
    style.font.name = ascii_font
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east)


def _apply_run_font(run, ascii_font: str, east_asia: Optional[str] = None) -> None:
    east = east_asia or ascii_font
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east)


def _set_doc_styles_english(document: Document) -> None:
    """
    English report typography: Cambria body (readable academic), Calibri headings (clear hierarchy),
    justified body text, relaxed vertical rhythm.
    """
    navy = RGBColor(0x14, 0x36, 0x5C)

    normal = document.styles["Normal"]
    _style_set_latin_font(normal, "Cambria")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.17
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_spec = [
        ("Heading 1", 17, Pt(18), Pt(10)),
        ("Heading 2", 13.5, Pt(14), Pt(7)),
        ("Heading 3", 12, Pt(12), Pt(6)),
    ]
    for name, size_pt, space_before, space_after in heading_spec:
        if name not in document.styles:
            continue
        st = document.styles[name]
        _style_set_latin_font(st, "Calibri")
        st.font.size = Pt(size_pt)
        st.font.bold = True
        st.font.color.rgb = navy
        st.paragraph_format.space_before = space_before
        st.paragraph_format.space_after = space_after
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        st.paragraph_format.line_spacing = 1.2

    for list_name in ("List Bullet", "List Number"):
        if list_name in document.styles:
            _style_set_latin_font(document.styles[list_name], "Cambria")


def _set_doc_styles_chinese(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)
    _ensure_style_eastasia(normal)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name not in document.styles:
            continue
        st = document.styles[name]
        st.font.name = "Calibri"
        _ensure_style_eastasia(st)
        st.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        if name == "Heading 1":
            st.font.size = Pt(16)
            st.font.bold = True
        elif name == "Heading 2":
            st.font.size = Pt(13)
            st.font.bold = True
        else:
            st.font.size = Pt(12)


def _set_doc_styles(document: Document, lang: Lang) -> None:
    if lang == "en":
        _set_doc_styles_english(document)
    else:
        _set_doc_styles_chinese(document)


def _add_cover(document: Document, lang: Lang, titles: Dict[str, str]) -> None:
    for _ in range(3):
        document.add_paragraph()

    def _east_asia(run, family: str = "Microsoft JhengHei") -> None:
        run.font.name = "Calibri"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
        rfonts.set(qn("w:eastAsia"), family)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sanitize_for_word(titles["title"]))
    r.bold = True
    if lang == "en":
        _apply_run_font(r, "Cambria")
        r.font.size = Pt(27)
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    else:
        r.font.size = Pt(26)
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        _east_asia(r)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sanitize_for_word(titles["subtitle"]))
    if lang == "en":
        _apply_run_font(r2, "Cambria")
        r2.font.size = Pt(14.5)
        r2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    else:
        r2.font.size = Pt(14)
        r2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        _east_asia(r2)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(sanitize_for_word(titles["tagline"]))
    if lang == "en":
        _apply_run_font(r3, "Cambria")
        r3.font.size = Pt(10.5)
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    else:
        r3.font.size = Pt(10)
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        _east_asia(r3)

    document.add_paragraph()
    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if lang == "zh":
        gen_date = dt.date.today().strftime("%Y 年 %m 月 %d 日")
        p4.add_run(sanitize_for_word(f"報告產製日期：{gen_date}")).font.size = Pt(11)
    else:
        gen_date = dt.date.today().strftime("%B %d, %Y")
        dr = p4.add_run(f"Generated: {gen_date}")
        _apply_run_font(dr, "Cambria")
        dr.font.size = Pt(11)
        dr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    document.add_page_break()


def _shade_row(table, row_idx: int, fill_hex: str) -> None:
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)


def _shade_header_row(table, fill_hex: str = "1E3A5F") -> None:
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def _add_table_from_rows(
    document: Document,
    headers: List[str],
    rows: List[List[str]],
    col_widths_cm: Optional[List[float]] = None,
) -> None:
    headers = [sanitize_for_word(h) for h in headers]
    rows = [[sanitize_for_word(str(c)) for c in row] for row in rows]
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
        if ri % 2 == 0:
            _shade_row(table, ri + 1, "F8FAFC")
    _shade_header_row(table)
    if col_widths_cm:
        for row in table.rows:
            for ci, w in enumerate(col_widths_cm):
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)
    document.add_paragraph()


def _load_csv_kv(path: Path) -> Dict[str, str]:
    if not path.exists():
        return {}
    import pandas as pd

    df = pd.read_csv(path, index_col=0)
    return {str(k): str(v) for k, v in df["value"].items()}


def _format_metrics_rows(
    meta: Dict[str, str], metrics: Dict[str, str], lang: Lang
) -> List[List[str]]:
    def fmt_pct_from_decimal(x: str, scale: float = 100.0) -> str:
        try:
            v = float(x)
            return f"{v * scale:.2f}%"
        except (ValueError, TypeError):
            return sanitize_for_word(x)

    def fmt_float(x: str, nd: int = 4) -> str:
        try:
            return f"{float(x):.{nd}f}"
        except (ValueError, TypeError):
            return sanitize_for_word(x)

    rows: List[List[str]] = []
    m = metrics

    if lang == "zh":
        labels = {
            "n": "樣本交易日數（測試集）",
            "cr_s": "策略累積報酬（淨值倍數 -1）",
            "cr_bh": "Buy and Hold 累積報酬",
            "cagr_s": "策略年化報酬（CAGR）",
            "cagr_bh": "B and H 年化報酬（CAGR）",
            "vol": "策略年化波動率",
            "sh_s": "策略 Sharpe（日頻、Rf 約 2% 年化）",
            "sh_bh": "B and H Sharpe",
            "mdd_s": "策略最大回撤（MDD）",
            "mdd_bh": "B and H 最大回撤（MDD）",
            "calmar": "卡瑪比率（年化除以絕對 MDD）",
            "win": "做多日勝率",
            "ridge": "Ridge 懲罰係數 lambda（驗證集選出）",
            "fee_b": "買進手續費率",
            "fee_s": "賣出手續費加證交稅",
        }
    else:
        labels = {
            "n": "Trading days (test set)",
            "cr_s": "Strategy cumulative return (wealth ratio - 1)",
            "cr_bh": "Buy and hold cumulative return",
            "cagr_s": "Strategy CAGR",
            "cagr_bh": "Buy and hold CAGR",
            "vol": "Strategy annualized volatility",
            "sh_s": "Strategy Sharpe (daily, Rf approx 2% p.a.)",
            "sh_bh": "Buy and hold Sharpe",
            "mdd_s": "Strategy max drawdown (MDD)",
            "mdd_bh": "Buy and hold MDD",
            "calmar": "Calmar ratio (est.)",
            "win": "Hit rate on long days",
            "ridge": "Ridge penalty lambda (from validation)",
            "fee_b": "Buy-side fee rate",
            "fee_s": "Sell fee plus transaction tax",
        }

    if "區間天數" in m:
        rows.append([labels["n"], m.get("區間天數", "N/A")])
    if "策略累積報酬" in m:
        cr = float(m["策略累積報酬"])
        rows.append(
            [
                labels["cr_s"],
                sanitize_for_word(f"{cr:.4f} (wealth approx {1 + cr:.2f}x)"),
            ]
        )
    if "B&H累積報酬" in m:
        cr = float(m["B&H累積報酬"])
        rows.append(
            [
                labels["cr_bh"],
                sanitize_for_word(f"{cr:.4f} (wealth approx {1 + cr:.2f}x)"),
            ]
        )
    if "策略年化報酬" in m:
        rows.append([labels["cagr_s"], fmt_pct_from_decimal(m["策略年化報酬"], 100)])
    if "B&H年化報酬" in m:
        rows.append([labels["cagr_bh"], fmt_pct_from_decimal(m["B&H年化報酬"], 100)])
    if "策略年化波動" in m:
        rows.append([labels["vol"], fmt_pct_from_decimal(m["策略年化波動"], 100)])
    if "策略Sharpe" in m:
        rows.append([labels["sh_s"], fmt_float(m["策略Sharpe"], 3)])
    if "B&H_Sharpe" in m:
        rows.append([labels["sh_bh"], fmt_float(m["B&H_Sharpe"], 3)])
    if "策略最大回撤%" in m:
        rows.append([labels["mdd_s"], f"{float(m['策略最大回撤%']):.2f}%"])
    if "B&H最大回撤%" in m:
        rows.append([labels["mdd_bh"], f"{float(m['B&H最大回撤%']):.2f}%"])
    if "卡瑪比率(估)" in m:
        rows.append([labels["calmar"], fmt_float(m["卡瑪比率(估)"], 2)])
    if "做多日勝率" in m:
        rows.append([labels["win"], fmt_pct_from_decimal(m["做多日勝率"], 100)])

    if meta.get("ridge_alpha"):
        rows.append([labels["ridge"], fmt_float(meta["ridge_alpha"], 6)])
    if meta.get("fee_buy"):
        rows.append([labels["fee_b"], fmt_pct_from_decimal(meta["fee_buy"], 100)])
    if meta.get("fee_sell_plus_tax"):
        rows.append(
            [labels["fee_s"], fmt_pct_from_decimal(meta["fee_sell_plus_tax"], 100)]
        )

    return rows


def _add_figure(
    document: Document,
    png_path: Path,
    caption: str,
    fig_no: int,
    lang: Lang,
    width_in: float = 6.85,
) -> None:
    caption = sanitize_for_word(caption)
    prefix = "Figure" if lang == "en" else "圖"
    sep = " " if lang == "en" else "　"
    if not png_path.exists():
        p = document.add_paragraph()
        miss = (
            f"[{prefix} {fig_no} missing: {png_path.name}; run run_overnight_signal.py]"
            if lang == "en"
            else f"〔{prefix} {fig_no} 缺檔：{png_path.name}，請先執行 run_overnight_signal.py〕"
        )
        mr = p.add_run(miss)
        mr.italic = True
        if lang == "en":
            _apply_run_font(mr, "Cambria")
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(png_path), width=Inches(width_in))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"{prefix} {fig_no}{sep}{caption}")
    cr.italic = True
    if lang == "en":
        _apply_run_font(cr, "Cambria")
        cr.font.size = Pt(10.5)
    else:
        cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def _add_figure_with_note(
    document: Document,
    png_path: Path,
    caption: str,
    fig_no: int,
    lang: Lang,
    interpretation: str,
    width_in: float = 6.85,
) -> None:
    _add_figure(document, png_path, caption, fig_no, lang, width_in=width_in)
    if interpretation.strip():
        document.add_paragraph(sanitize_for_word(interpretation))
    document.add_paragraph()


def _paragraph(document: Document, text: str) -> None:
    document.add_paragraph(sanitize_for_word(text))


def _method_usage_rows(lang: Lang) -> List[List[str]]:
    if lang == "zh":
        return [
            [
                "Ridge 迴歸（L2 正則化）",
                "降低過擬合與共線性風險，讓係數更穩定，提升樣本外泛化能力。",
            ],
            [
                "訓練／驗證／測試切分",
                "隔離調參與最終評估，避免資料洩漏與過度樂觀績效。",
            ],
            ["對數網格搜尋 lambda", "系統化尋找正則化強度，在偏差與變異之間取得平衡。"],
            [
                "做多或空手規則（long or flat）",
                "將連續預測轉為可執行訊號，降低雙邊摩擦與放空限制影響。",
            ],
            [
                "交易成本建模（手續費、證交稅）",
                "將理論報酬折減為較接近實務可達成的淨報酬。",
            ],
            [
                "資產淨值曲線（Equity Curve）",
                "觀察策略資本隨時間的成長路徑與複利效果。",
            ],
            [
                "最大回撤（MDD / Underwater）",
                "衡量從高點回落的最壞跌幅，評估壓力承受能力。",
            ],
            ["滾動 Sharpe 比率", "追蹤風險調整報酬是否持續，辨識績效是否衰退。"],
            [
                "年度 IC 與滾動 IC",
                "檢查預測訊號與實現報酬的關聯強度與穩定性（alpha 衰減）。",
            ],
            ["月報酬熱力圖", "辨識報酬的時間分布與市場狀態敏感性。"],
            ["Buy and Hold 基準比較", "判斷策略是否提供超額報酬與較佳風險控制。"],
        ]
    return [
        [
            "Ridge regression (L2 regularization)",
            "Controls overfitting and multicollinearity, producing more stable coefficients and better out-of-sample robustness.",
        ],
        [
            "Train / validation / test split",
            "Separates tuning from final evaluation to reduce leakage and optimistic bias.",
        ],
        [
            "Log-grid search for lambda",
            "Systematically selects regularization strength to balance bias and variance.",
        ],
        [
            "Long-or-flat signal mapping",
            "Converts continuous forecasts into executable positions while reducing shorting and friction constraints.",
        ],
        [
            "Transaction cost modeling (fees, tax)",
            "Transforms gross returns into net returns closer to realistic implementation.",
        ],
        [
            "Equity curve",
            "Shows capital growth path and compounding behavior over time.",
        ],
        [
            "Maximum drawdown (MDD / underwater)",
            "Measures worst peak-to-trough loss and downside tolerance.",
        ],
        [
            "Rolling Sharpe ratio",
            "Monitors persistence of risk-adjusted performance through time.",
        ],
        [
            "Annual IC and rolling IC",
            "Evaluates predictive signal relevance and stability, including potential alpha decay.",
        ],
        [
            "Monthly return heatmap",
            "Visualizes return clustering and regime sensitivity by calendar month.",
        ],
        [
            "Buy-and-hold benchmark comparison",
            "Checks whether the strategy adds alpha and improves risk control versus passive exposure.",
        ],
    ]
