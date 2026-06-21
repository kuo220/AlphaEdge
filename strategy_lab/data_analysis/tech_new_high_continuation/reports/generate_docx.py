#!/usr/bin/env python3
"""
將 tech_new_high_continuation 分析 output/ CSV 彙整為精簡 Word 報告（.docx）。

使用方式（專案根目錄）：
    .venv/bin/python strategy_lab/data_analysis/tech_new_high_continuation/reports/generate_docx.py

若尚未有 CSV，請先執行：
    .venv/bin/python strategy_lab/data_analysis/tech_new_high_continuation/run.py
"""

from __future__ import annotations

import datetime as dt
import sys
from pathlib import Path
from typing import List, Optional

_PROJECT_ROOT = Path(__file__).resolve().parents[4]
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as e:
    raise SystemExit(
        "請先安裝 python-docx：pip install python-docx\n" + str(e)
    ) from e

import pandas as pd

from strategy_lab.data_analysis.tech_new_high_continuation.analysis import (  # noqa: E402
    END_DATE,
    HORIZONS,
    MIN_EVENT_COUNT,
    START_DATE,
)

_ANALYSIS_DIR = Path(__file__).resolve().parent.parent
_OUTPUT_DIR = _ANALYSIS_DIR / "output"
_REPORTS_DIR = Path(__file__).resolve().parent
_DOCX_PATH = _REPORTS_DIR / "Tech_NewHigh_Continuation_Report.docx"

TOP_N = 20
BOTTOM_N = 15
MIN_INDUSTRY_EVENTS = 50


def _ensure_eastasia(run, family: str = "Microsoft JhengHei") -> None:
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), family)


def _set_doc_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        if name not in document.styles:
            continue
        st = document.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        rpr = st._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def _shade_row(table, row_idx: int, fill_hex: str) -> None:
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)


def _shade_header_row(table, fill_hex: str = "1E3A5F") -> None:
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def _add_table(
    document: Document,
    headers: List[str],
    rows: List[List[str]],
    col_widths_cm: Optional[List[float]] = None,
) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
        if ri % 2 == 0:
            _shade_row(table, ri + 1, "F8FAFC")
    _shade_header_row(table)
    if col_widths_cm:
        for row in table.rows:
            for ci, w in enumerate(col_widths_cm):
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)
    document.add_paragraph()


def _pct(x: float) -> str:
    return f"{x * 100:.1f}%"


def _load_data() -> tuple[pd.DataFrame, pd.DataFrame]:
    summary_path = _OUTPUT_DIR / "tech_new_high_prob_summary.csv"
    by_stock_path = _OUTPUT_DIR / "tech_new_high_prob_by_stock.csv"
    if not summary_path.exists() or not by_stock_path.exists():
        raise FileNotFoundError(
            "找不到 output CSV，請先執行 run.py：\n"
            "  .venv/bin/python strategy_lab/data_analysis/tech_new_high_continuation/run.py"
        )
    return pd.read_csv(summary_path), pd.read_csv(by_stock_path)


def _stock_table_rows(df: pd.DataFrame) -> List[List[str]]:
    rows: List[List[str]] = []
    for _, r in df.iterrows():
        rows.append(
            [
                str(int(r["stock_id"])).zfill(4),
                str(r["stock_name"]),
                str(r["industry"]),
                str(int(r["event_count"])),
                _pct(r["prob_5d"]),
                _pct(r["prob_10d"]),
                _pct(r["prob_20d"]),
            ]
        )
    return rows


def build_report() -> Path:
    summary, by_stock = _load_data()

    non = summary[summary["event_mode"] == "non_overlap"].copy()
    all_ev = summary[summary["event_mode"] == "all_events"].copy()
    overall_non = non[non["scope"] == "all_tech"].iloc[0]
    overall_all = all_ev[all_ev["scope"] == "all_tech"].iloc[0]

    by_ind_non = non[non["scope"].str.startswith("by_industry:")].copy()
    by_ind_non["industry"] = by_ind_non["scope"].str.replace("by_industry:", "", regex=False)
    by_ind_non = by_ind_non.sort_values("prob_20d", ascending=False)

    all_stocks = by_stock[
        (by_stock["event_mode"] == "all_events") & (~by_stock["low_sample"])
    ].copy()
    top_stocks = all_stocks.nlargest(TOP_N, "prob_20d")
    bottom_stocks = all_stocks.nsmallest(BOTTOM_N, "prob_20d")

    focus_ids = ["2330", "2454", "2317", "2382", "2303"]
    focus = by_stock[
        (by_stock["event_mode"] == "all_events")
        & (by_stock["stock_id"].astype(str).isin(focus_ids))
    ].sort_values("stock_id")

    ind_sufficient = by_ind_non[by_ind_non["event_count"] >= MIN_INDUSTRY_EVENTS]
    best_ind = ind_sufficient.iloc[0] if not ind_sufficient.empty else by_ind_non.iloc[0]
    worst_ind = (
        ind_sufficient.iloc[-1] if not ind_sufficient.empty else by_ind_non.iloc[-1]
    )

    doc = Document()
    _set_doc_styles(doc)

    for _ in range(2):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("台股科技業創高後續上行機率分析")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    _ensure_eastasia(tr)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("歷史新高事件之統計摘要報告")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    _ensure_eastasia(sr)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"產製日期：{dt.date.today().strftime('%Y 年 %m 月 %d 日')}")
    dr.font.size = Pt(10)
    _ensure_eastasia(dr)

    doc.add_page_break()

    doc.add_heading("一、研究摘要", level=1)
    bullets = [
        f"樣本期間 {START_DATE} ~ {END_DATE}，涵蓋台股科技業普通股 "
        f"{len(all_stocks) + int(by_stock[by_stock['event_mode'] == 'all_events']['low_sample'].sum())} 檔"
        f"（有效樣本 ≥ {MIN_EVENT_COUNT} 次創高者 {len(all_stocks)} 檔）。",
        "創高定義：最高價突破歷史新高；持續上行：創高後 N 日收盤價高於創高日收盤價。",
        f"主結果（non_overlap，20 日 cooldown）：整體 20 日持續上行機率 "
        f"{_pct(overall_non['prob_20d'])}（{int(overall_non['event_count']):,} 次事件）。",
        f"5 / 10 / 20 日機率分別為 {_pct(overall_non['prob_5d'])}、"
        f"{_pct(overall_non['prob_10d'])}、{_pct(overall_non['prob_20d'])}，"
        "略高於隨機猜測（50%），動能延續效果有限。",
        f"分產業（non_overlap，樣本 ≥ {MIN_INDUSTRY_EVENTS}）：20 日機率最高為 "
        f"{best_ind['industry']}（{_pct(best_ind['prob_20d'])}），"
        f"最低為 {worst_ind['industry']}（{_pct(worst_ind['prob_20d'])}）。",
        f"逐檔（all_events，樣本 ≥ {MIN_EVENT_COUNT}）：20 日機率中位數 "
        f"{_pct(all_stocks['prob_20d'].median())}，"
        f"個股差異大（最高 {int(top_stocks.iloc[0]['stock_id']):04d} "
        f"{top_stocks.iloc[0]['stock_name']} {_pct(top_stocks.iloc[0]['prob_20d'])}；"
        f"最低 {int(bottom_stocks.iloc[0]['stock_id']):04d} "
        f"{bottom_stocks.iloc[0]['stock_name']} {_pct(bottom_stocks.iloc[0]['prob_20d'])}）。",
    ]
    for text in bullets:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("二、方法與資料", level=1)
    method_rows = [
        ["標的", "台股 11 類科技／電子產業普通股"],
        ["資料來源", "SQLite price 表 + taiwan_stock_info"],
        ["樣本期間", f"{START_DATE} ~ {END_DATE}"],
        ["創高", "當日最高價 > 截至前一交易日之歷史最高價"],
        ["持續上行", "創高後第 N 日收盤價 > 創高日收盤價"],
        ["持有期 N", "5 / 10 / 20 個交易日"],
        ["事件去重", "non_overlap：創高後 20 日內不重複計入（主報告）"],
        ["對照組", "all_events：所有創高日皆計入"],
        ["有效個股門檻", f"創高事件數 ≥ {MIN_EVENT_COUNT} 次"],
    ]
    _add_table(doc, ["項目", "說明"], method_rows, [4.0, 12.0])

    doc.add_heading("三、整體統計結果", level=1)
    doc.add_paragraph(
        "下表為科技業池化後的持續上行機率。"
        "建議以 non_overlap 解讀（事件獨立性較佳）；all_events 樣本較大但事件高度相關。"
    )
    overall_rows = [
        [
            "non_overlap（主）",
            f"{int(overall_non['event_count']):,}",
            _pct(overall_non["prob_5d"]),
            _pct(overall_non["prob_10d"]),
            _pct(overall_non["prob_20d"]),
        ],
        [
            "all_events（對照）",
            f"{int(overall_all['event_count']):,}",
            _pct(overall_all["prob_5d"]),
            _pct(overall_all["prob_10d"]),
            _pct(overall_all["prob_20d"]),
        ],
    ]
    _add_table(
        doc,
        ["統計口徑", "事件數", "5 日機率", "10 日機率", "20 日機率"],
        overall_rows,
        [4.5, 2.5, 2.5, 2.5, 2.5],
    )

    doc.add_heading("四、分產業統計（non_overlap）", level=1)
    ind_rows: List[List[str]] = []
    for _, r in by_ind_non.iterrows():
        note = ""
        if r["event_count"] < MIN_INDUSTRY_EVENTS:
            note = " *"
        ind_rows.append(
            [
                str(r["industry"]) + note,
                str(int(r["event_count"])),
                _pct(r["prob_5d"]),
                _pct(r["prob_10d"]),
                _pct(r["prob_20d"]),
            ]
        )
    _add_table(
        doc,
        ["產業", "事件數", "5 日", "10 日", "20 日"],
        ind_rows,
        [4.5, 2.0, 2.0, 2.0, 2.0],
    )
    doc.add_paragraph(f"* 事件數 < {MIN_INDUSTRY_EVENTS}，解讀時請謹慎。")

    doc.add_heading(f"五、逐檔表現 — 20 日機率 Top {TOP_N}", level=1)
    doc.add_paragraph(
        f"以下為 all_events 口徑、創高次數 ≥ {MIN_EVENT_COUNT} 的個股，"
        "依 20 日持續上行機率由高到低排序。"
    )
    _add_table(
        doc,
        ["代號", "名稱", "產業", "創高次數", "5 日", "10 日", "20 日"],
        _stock_table_rows(top_stocks),
        [1.5, 2.5, 3.0, 2.0, 1.8, 1.8, 1.8],
    )

    doc.add_heading(f"六、逐檔表現 — 20 日機率 Bottom {BOTTOM_N}", level=1)
    _add_table(
        doc,
        ["代號", "名稱", "產業", "創高次數", "5 日", "10 日", "20 日"],
        _stock_table_rows(bottom_stocks),
        [1.5, 2.5, 3.0, 2.0, 1.8, 1.8, 1.8],
    )

    doc.add_heading("七、代表性個股", level=1)
    doc.add_paragraph("科技權值股之創高後持續上行機率（all_events）：")
    _add_table(
        doc,
        ["代號", "名稱", "產業", "創高次數", "5 日", "10 日", "20 日"],
        _stock_table_rows(focus),
        [1.5, 2.5, 3.0, 2.0, 1.8, 1.8, 1.8],
    )

    doc.add_heading("八、結論與限制", level=1)
    conclusions = [
        "整體而言，科技股創歷史新高後 20 日收盤仍高於創高日的機率約 50%，"
        "並未呈現強烈動能延續，不宜單純以「創高」作為追價依據。",
        "個股與產業間差異顯著：部分標的長期創高後續航率偏高，"
        "亦有標的創高後快速回落，需搭配個股樣本數一併評估。",
        "本分析使用未還原價，除權息可能造成假創高；產業分類為最新快照，"
        "未反映歷史產業變更。",
        "完整事件明細與逐檔統計見 output/ 目錄下 CSV 檔。",
    ]
    for text in conclusions:
        doc.add_paragraph(text, style="List Bullet")

    _REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(_DOCX_PATH)
    return _DOCX_PATH


def main() -> Path:
    path = build_report()
    print(f"報告已寫入：{path}")
    return path


if __name__ == "__main__":
    main()
