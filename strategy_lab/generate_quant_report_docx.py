#!/usr/bin/env python3
"""
將 strategy_lab/output 之圖表與 CSV 彙整為 Word 量化報告（.docx），支援中文／英文。

使用方式（專案根目錄）：
    .venv/bin/python strategy_lab/generate_quant_report_docx.py
    .venv/bin/python strategy_lab/generate_quant_report_docx.py --lang en
    .venv/bin/python strategy_lab/generate_quant_report_docx.py --lang both

若尚未有圖表，請先執行：.venv/bin/python strategy_lab/run_overnight_signal.py
"""

from __future__ import annotations

import argparse
import datetime as dt
import re
from pathlib import Path
from typing import Dict, List, Literal, Optional, Tuple

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as e:
    raise SystemExit(
        "請先安裝 python-docx：pip install python-docx\n" + str(e)
    ) from e

_OUTPUT = Path(__file__).resolve().parent / "output"
_DOCX_ZH = Path(__file__).resolve().parent / "TSMC_OvernightSignal_Quant_Report.docx"
_DOCX_EN = Path(__file__).resolve().parent / "TSMC_OvernightSignal_Quant_Report_EN.docx"

Lang = Literal["zh", "en"]


def sanitize_for_word(text: str) -> str:
    """
    過濾易在 Word 造成顯示異常或與 LaTeX 混淆的字元。
    - 移除 $（避免被誤認為數學模式）
    - 替換全形／特殊連字與異常空白
    - 將 Unicode 減號統一為 ASCII 連字號（表格與數字可讀性）
    """
    if not text:
        return text
    s = text.replace("$", "")
    s = s.replace("\u00a0", " ")  # nbsp
    s = s.replace("\u2028", " ")
    s = s.replace("\u2029", " ")
    s = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", s)  # zero-width
    s = s.replace("\u2212", "-").replace("\u2013", "-").replace("\u2014", "-")
    s = s.replace("§", "")
    return s.strip()


# ---------------------------------------------------------------------------
# 標題與用語（統一使用隔夜訊號／跨市場領先）
# ---------------------------------------------------------------------------
TITLES_ZH = {
    "title": "台積電跨市場定價與隔夜訊號研究",
    "subtitle": "ADR、費城半導體與匯率之領先資訊：Ridge 預測與歷史回測",
    "tagline": "量化回測報告 · AlphaEdge Strategy Workspace",
}

TITLES_EN = {
    "title": "TSMC Cross-Market Pricing and Overnight Information",
    "subtitle": "Lead Signals from ADR, SOX, and FX: Ridge Forecast and Backtest",
    "tagline": "Quantitative Backtest Report · AlphaEdge Strategy Workspace",
}


def _ensure_style_eastasia(style, east: str = "Microsoft JhengHei") -> None:
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), east)


def _set_doc_styles(document: Document, lang: Lang) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)
    if lang == "zh":
        _ensure_style_eastasia(normal)
    else:
        _ensure_style_eastasia(normal, east="Calibri")

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name not in document.styles:
            continue
        st = document.styles[name]
        st.font.name = "Calibri"
        if lang == "zh":
            _ensure_style_eastasia(st)
        else:
            _ensure_style_eastasia(st, east="Calibri")
        st.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        if name == "Heading 1":
            st.font.size = Pt(16)
            st.font.bold = True
        elif name == "Heading 2":
            st.font.size = Pt(13)
            st.font.bold = True
        else:
            st.font.size = Pt(12)


def _add_cover(document: Document, lang: Lang, titles: Dict[str, str]) -> None:
    for _ in range(3):
        document.add_paragraph()

    def _east_asia(run, family: str = "Microsoft JhengHei") -> None:
        run.font.name = "Calibri"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
        rfonts.set(qn("w:eastAsia"), family)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sanitize_for_word(titles["title"]))
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    if lang == "zh":
        _east_asia(r)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sanitize_for_word(titles["subtitle"]))
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    if lang == "zh":
        _east_asia(r2)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(sanitize_for_word(titles["tagline"]))
    r3.font.size = Pt(10)
    r3.italic = True
    r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    if lang == "zh":
        _east_asia(r3)

    document.add_paragraph()
    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if lang == "zh":
        gen_date = dt.date.today().strftime("%Y 年 %m 月 %d 日")
        p4.add_run(sanitize_for_word(f"報告產製日期：{gen_date}")).font.size = Pt(11)
    else:
        gen_date = dt.date.today().strftime("%B %d, %Y")
        p4.add_run(f"Generated: {gen_date}").font.size = Pt(11)

    document.add_page_break()


def _shade_row(table, row_idx: int, fill_hex: str) -> None:
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)


def _shade_header_row(table, fill_hex: str = "1E3A5F") -> None:
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def _add_table_from_rows(
    document: Document,
    headers: List[str],
    rows: List[List[str]],
    col_widths_cm: Optional[List[float]] = None,
) -> None:
    headers = [sanitize_for_word(h) for h in headers]
    rows = [[sanitize_for_word(str(c)) for c in row] for row in rows]
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
        if ri % 2 == 0:
            _shade_row(table, ri + 1, "F8FAFC")
    _shade_header_row(table)
    if col_widths_cm:
        for row in table.rows:
            for ci, w in enumerate(col_widths_cm):
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)
    document.add_paragraph()


def _load_csv_kv(path: Path) -> Dict[str, str]:
    if not path.exists():
        return {}
    import pandas as pd

    df = pd.read_csv(path, index_col=0)
    return {str(k): str(v) for k, v in df["value"].items()}


def _format_metrics_rows(
    meta: Dict[str, str], metrics: Dict[str, str], lang: Lang
) -> List[List[str]]:
    def fmt_pct_from_decimal(x: str, scale: float = 100.0) -> str:
        try:
            v = float(x)
            return f"{v * scale:.2f}%"
        except (ValueError, TypeError):
            return sanitize_for_word(x)

    def fmt_float(x: str, nd: int = 4) -> str:
        try:
            return f"{float(x):.{nd}f}"
        except (ValueError, TypeError):
            return sanitize_for_word(x)

    rows: List[List[str]] = []
    m = metrics

    if lang == "zh":
        labels = {
            "n": "樣本交易日數（測試集）",
            "cr_s": "策略累積報酬（淨值倍數 -1）",
            "cr_bh": "Buy and Hold 累積報酬",
            "cagr_s": "策略年化報酬（CAGR）",
            "cagr_bh": "B and H 年化報酬（CAGR）",
            "vol": "策略年化波動率",
            "sh_s": "策略 Sharpe（日頻、Rf 約 2% 年化）",
            "sh_bh": "B and H Sharpe",
            "mdd_s": "策略最大回撤（MDD）",
            "mdd_bh": "B and H 最大回撤（MDD）",
            "calmar": "卡瑪比率（年化除以絕對 MDD）",
            "win": "做多日勝率",
            "ridge": "Ridge 懲罰係數 lambda（驗證集選出）",
            "fee_b": "買進手續費率",
            "fee_s": "賣出手續費加證交稅",
        }
    else:
        labels = {
            "n": "Trading days (test set)",
            "cr_s": "Strategy cumulative return (wealth ratio - 1)",
            "cr_bh": "Buy and hold cumulative return",
            "cagr_s": "Strategy CAGR",
            "cagr_bh": "Buy and hold CAGR",
            "vol": "Strategy annualized volatility",
            "sh_s": "Strategy Sharpe (daily, Rf approx 2% p.a.)",
            "sh_bh": "Buy and hold Sharpe",
            "mdd_s": "Strategy max drawdown (MDD)",
            "mdd_bh": "Buy and hold MDD",
            "calmar": "Calmar ratio (est.)",
            "win": "Hit rate on long days",
            "ridge": "Ridge penalty lambda (from validation)",
            "fee_b": "Buy-side fee rate",
            "fee_s": "Sell fee plus transaction tax",
        }

    if "區間天數" in m:
        rows.append([labels["n"], m.get("區間天數", "N/A")])
    if "策略累積報酬" in m:
        cr = float(m["策略累積報酬"])
        rows.append(
            [
                labels["cr_s"],
                sanitize_for_word(f"{cr:.4f} (wealth approx {1+cr:.2f}x)"),
            ]
        )
    if "B&H累積報酬" in m:
        cr = float(m["B&H累積報酬"])
        rows.append(
            [
                labels["cr_bh"],
                sanitize_for_word(f"{cr:.4f} (wealth approx {1+cr:.2f}x)"),
            ]
        )
    if "策略年化報酬" in m:
        rows.append([labels["cagr_s"], fmt_pct_from_decimal(m["策略年化報酬"], 100)])
    if "B&H年化報酬" in m:
        rows.append([labels["cagr_bh"], fmt_pct_from_decimal(m["B&H年化報酬"], 100)])
    if "策略年化波動" in m:
        rows.append([labels["vol"], fmt_pct_from_decimal(m["策略年化波動"], 100)])
    if "策略Sharpe" in m:
        rows.append([labels["sh_s"], fmt_float(m["策略Sharpe"], 3)])
    if "B&H_Sharpe" in m:
        rows.append([labels["sh_bh"], fmt_float(m["B&H_Sharpe"], 3)])
    if "策略最大回撤%" in m:
        rows.append([labels["mdd_s"], f"{float(m['策略最大回撤%']):.2f}%"])
    if "B&H最大回撤%" in m:
        rows.append([labels["mdd_bh"], f"{float(m['B&H最大回撤%']):.2f}%"])
    if "卡瑪比率(估)" in m:
        rows.append([labels["calmar"], fmt_float(m["卡瑪比率(估)"], 2)])
    if "做多日勝率" in m:
        rows.append([labels["win"], fmt_pct_from_decimal(m["做多日勝率"], 100)])

    if meta.get("ridge_alpha"):
        rows.append([labels["ridge"], fmt_float(meta["ridge_alpha"], 6)])
    if meta.get("fee_buy"):
        rows.append([labels["fee_b"], fmt_pct_from_decimal(meta["fee_buy"], 100)])
    if meta.get("fee_sell_plus_tax"):
        rows.append([labels["fee_s"], fmt_pct_from_decimal(meta["fee_sell_plus_tax"], 100)])

    return rows


def _add_figure(
    document: Document,
    png_path: Path,
    caption: str,
    fig_no: int,
    lang: Lang,
    width_in: float = 6.2,
) -> None:
    caption = sanitize_for_word(caption)
    prefix = "Figure" if lang == "en" else "圖"
    sep = " " if lang == "en" else "　"
    if not png_path.exists():
        p = document.add_paragraph()
        miss = (
            f"[{prefix} {fig_no} missing: {png_path.name}; run run_overnight_signal.py]"
            if lang == "en"
            else f"〔{prefix} {fig_no} 缺檔：{png_path.name}，請先執行 run_overnight_signal.py〕"
        )
        p.add_run(miss).italic = True
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(png_path), width=Inches(width_in))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"{prefix} {fig_no}{sep}{caption}")
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    document.add_paragraph()


def _paragraph(document: Document, text: str) -> None:
    document.add_paragraph(sanitize_for_word(text))


def _method_usage_rows(lang: Lang) -> List[List[str]]:
    if lang == "zh":
        return [
            ["Ridge 迴歸（L2 正則化）", "降低過擬合與共線性風險，讓係數更穩定，提升樣本外泛化能力。"],
            ["訓練／驗證／測試切分", "隔離調參與最終評估，避免資料洩漏與過度樂觀績效。"],
            ["對數網格搜尋 lambda", "系統化尋找正則化強度，在偏差與變異之間取得平衡。"],
            ["做多或空手規則（long or flat）", "將連續預測轉為可執行訊號，降低雙邊摩擦與放空限制影響。"],
            ["交易成本建模（手續費、證交稅）", "將理論報酬折減為較接近實務可達成的淨報酬。"],
            ["資產淨值曲線（Equity Curve）", "觀察策略資本隨時間的成長路徑與複利效果。"],
            ["最大回撤（MDD / Underwater）", "衡量從高點回落的最壞跌幅，評估壓力承受能力。"],
            ["滾動 Sharpe 比率", "追蹤風險調整報酬是否持續，辨識績效是否衰退。"],
            ["年度 IC 與滾動 IC", "檢查預測訊號與實現報酬的關聯強度與穩定性（alpha 衰減）。"],
            ["月報酬熱力圖", "辨識報酬的時間分布與市場狀態敏感性。"],
            ["Buy and Hold 基準比較", "判斷策略是否提供超額報酬與較佳風險控制。"],
        ]
    return [
        ["Ridge regression (L2 regularization)", "Controls overfitting and multicollinearity, producing more stable coefficients and better out-of-sample robustness."],
        ["Train / validation / test split", "Separates tuning from final evaluation to reduce leakage and optimistic bias."],
        ["Log-grid search for lambda", "Systematically selects regularization strength to balance bias and variance."],
        ["Long-or-flat signal mapping", "Converts continuous forecasts into executable positions while reducing shorting and friction constraints."],
        ["Transaction cost modeling (fees, tax)", "Transforms gross returns into net returns closer to realistic implementation."],
        ["Equity curve", "Shows capital growth path and compounding behavior over time."],
        ["Maximum drawdown (MDD / underwater)", "Measures worst peak-to-trough loss and downside tolerance."],
        ["Rolling Sharpe ratio", "Monitors persistence of risk-adjusted performance through time."],
        ["Annual IC and rolling IC", "Evaluates predictive signal relevance and stability, including potential alpha decay."],
        ["Monthly return heatmap", "Visualizes return clustering and regime sensitivity by calendar month."],
        ["Buy-and-hold benchmark comparison", "Checks whether the strategy adds alpha and improves risk control versus passive exposure."],
    ]


def build_report(
    lang: Lang,
    output_path: Optional[Path] = None,
) -> Path:
    titles = TITLES_EN if lang == "en" else TITLES_ZH
    out_doc = output_path or (_DOCX_EN if lang == "en" else _DOCX_ZH)
    meta = _load_csv_kv(_OUTPUT / "run_meta.csv")
    metrics_raw = _load_csv_kv(_OUTPUT / "metrics_summary.csv")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _set_doc_styles(doc, lang)
    _add_cover(doc, lang, titles)

    if lang == "zh":
        doc.add_heading("摘要", level=1)
        _paragraph(
            doc,
            "本研究探討台積電美國存託憑證與台股現貨之跨市場資訊傳遞：在美國市場與匯率已於台股開盤前結束交易的前提下，"
            "以台積電 ADR（TSM）、費城半導體指數（SOX）與台幣兌美元（TWD=X）之報酬作為領先特徵，"
            "以 Ridge 迴歸預測台股 2330.TW 當日報酬，並採做多或空手之日頻簡化回測。"
            "資料來自 yfinance；指標命名對齊 AlphaEdge 回測慣例。",
        )
        note = doc.add_paragraph()
        note.add_run(
            sanitize_for_word(
                "聲明：本文件僅供學術與課程專題展示，不構成投資建議；實務交易含流動性、稅費、法規與執行落差。"
            )
        ).italic = True

        doc.add_heading("一、研究設計與樣本切分", level=1)
        _paragraph(
            doc,
            "訓練集：至 2020-12-31；驗證集：2021 年度（用於 Ridge 懲罰係數 lambda 之網格搜尋）；"
            "測試集：2022-01-01 至 2026-04-25。"
            "特徵對齊：每一台股交易日 t，取嚴格早於 t 之最後一個美國交易日之報酬，以避免前視偏差。",
        )
        _add_table_from_rows(
            doc,
            ["資料集欄位", "內容"],
            [
                ["回測總區間", "2020-01-01 ~ 2026-04-25"],
                ["訓練集", "2020-01-01 ~ 2020-12-31"],
                ["驗證集", "2021-01-01 ~ 2021-12-31"],
                ["測試集", "2022-01-01 ~ 2026-04-25"],
                ["資料來源", "yfinance（日線、還原權息）"],
                ["標的與特徵", "2330.TW（目標）；TSM、^SOX、TWD=X（特徵）"],
            ],
            col_widths_cm=[4.5, 12],
        )

        doc.add_heading("二、策略邏輯說明", level=1)
        _paragraph(
            doc,
            "（1）經濟直覺：台股開盤前，美國已歷經完整交易時段，TSM 與半導體板塊（SOX）及台幣走勢已反映隔夜全球資訊，"
            "理論上可作為隔日台股報酬之領先指標。",
        )
        _paragraph(
            doc,
            "（2）時間對齊：對每個 2330.TW 有成交之日 d，特徵取日曆上早於 d 之最後美國交易日之報酬，"
            "確保在台股開盤時點可觀測。",
        )
        _paragraph(
            doc,
            "（3）特徵與標的：自變數為該美國交易日之 TSM、SOX、TWD=X 日報酬率；"
            "應變數為同日台股 2330.TW 之收盤對前收報酬率（日頻持有期近似）。",
        )
        _paragraph(
            doc,
            "（4）估計：以含截距之 Ridge 迴歸估計線性關係，懲罰係數 lambda 在對數網格上搜尋，並以 2021 年驗證集均方誤差決定。",
        )
        _paragraph(
            doc,
            "（5）交易規則：若模型預測之當日報酬大於 0，則以全額資金持有 2330 現貨多單；否則空手（報酬率視為 0，不計融資成本）。",
        )
        _paragraph(
            doc,
            "（6）成本：加碼多單時扣買進手續費；減碼或平倉時扣賣出手續費與證交稅，與期末報告假設一致。",
        )
        _paragraph(
            doc,
            "（7）買賣時點細化：當 pred_t 大於 0 且前一日空手時，當日視為買進建倉；當 pred_t 小於等於 0 且前一日持有時，"
            "當日視為賣出平倉；若連續兩天同為持有或同為空手，則不發生換手成本。",
        )
        _paragraph(
            doc,
            "（8）使用 Ridge 的目的：在 TSM、SOX 等高相關特徵下降低共線性衝擊，並以 L2 正則化抑制過擬合，"
            "提升樣本外預測穩定度，同時維持線性模型的可解釋性。",
        )

        doc.add_heading("三、變數與模型", level=1)
        _add_table_from_rows(
            doc,
            ["資料代號（yfinance）", "角色"],
            [
                ["TSM", "台積電 ADR，美國市場前一日報酬"],
                ["2330.TW", "台股報酬（標的與 Buy and Hold 基準）"],
                ["^SOX", "費城半導體指數，產業情緒與 beta 代理"],
                ["TWD=X", "台幣兌美元即期（NDF 之簡化代理）"],
            ],
            col_widths_cm=[4.5, 12],
        )
        _paragraph(
            doc,
            "模型為 Ridge 迴歸，L2 僅施加於斜率；lambda 由驗證集決定後，於訓練加驗證樣本重估係數，再於測試集評估。",
        )

        doc.add_heading("四、方法與用途說明", level=1)
        _add_table_from_rows(
            doc,
            ["方法", "用途說明"],
            _method_usage_rows("zh"),
            col_widths_cm=[6, 10.5],
        )

        doc.add_heading("五、測試集績效彙總", level=1)
        perf_rows = _format_metrics_rows(meta, metrics_raw, "zh")
        if perf_rows:
            _add_table_from_rows(doc, ["指標", "數值"], perf_rows, col_widths_cm=[5.5, 11])
        else:
            _paragraph(
                doc,
                "（尚未偵測到 output/metrics_summary.csv，請先於專案根目錄執行 .venv/bin/python strategy_lab/run_overnight_signal.py）",
            )

        doc.add_heading("六、圖表與視覺化分析", level=1)
        _paragraph(
            doc,
            "以下圖表由回測程式輸出（PNG），依序為淨值、回撤、滾動夏普、月報酬熱圖、年度 IC、滾動 IC。",
        )
        figures_zh: List[Tuple[str, str]] = [
            ("equity_curve.png", "資產淨值曲線：本策略與 Buy and Hold 2330.TW"),
            ("mdd_underwater.png", "最大回撤（水下曲線）：策略與基準"),
            ("rolling_sharpe.png", "滾動夏普比率（約 63 交易日，年化）"),
            ("monthly_returns_heatmap.png", "月報酬率熱力圖（百分比）"),
            ("ic_by_year.png", "預測與實現報酬之年度資訊係數 IC"),
            ("rolling_ic.png", "滾動 IC（126 交易日窗口）"),
        ]
        for i, (fname, cap) in enumerate(figures_zh, start=1):
            _add_figure(doc, _OUTPUT / fname, cap, i, "zh")

        doc.add_heading("七、風險因素與限制", level=1)
        for t in [
            "線性 Ridge 可能無法捕捉非線性與結構斷點；未實作梯度提升樹與線上重訓。",
            "匯率使用即期代理，非避險後之淨曝險。",
            "未模擬漲跌停、開盤撮合、滑價與期貨保證金與換月。",
            "僅做多或空手，未涵蓋融券或指數期貨。",
            "驗證集參與 lambda 選擇可能帶來輕度資料探勘。",
        ]:
            doc.add_paragraph(sanitize_for_word(t), style="List Bullet")

        doc.add_heading("八、重現方式", level=1)
        _paragraph(
            doc,
            "於 AlphaEdge 專案根目錄執行 .venv/bin/python strategy_lab/run_overnight_signal.py 產出圖表與 CSV；"
            "再執行 .venv/bin/python strategy_lab/generate_quant_report_docx.py 產生本報告。",
        )

        doc.add_heading("參考資料", level=1)
        doc.add_paragraph(
            sanitize_for_word("yfinance；AlphaEdge core/backtest/README.md；課程期末報告原始稿（成本與樣本假設）。"),
            style="List Bullet",
        )

    else:
        doc.add_heading("Executive Summary", level=1)
        _paragraph(
            doc,
            "This study examines cross-market information flow for TSMC between the U.S. ADR and the Taiwan listing. "
            "Because the U.S. session ends before the Taiwan cash market opens, returns of TSMC ADR (TSM), the PHLX "
            "Semiconductor Index (SOX), and USD/TWD (TWD=X) are used as lead features available prior to the Taiwan open. "
            "A Ridge regression forecasts the same-day close-to-close return of 2330.TW; the trading rule is long the "
            "stock when the forecast is positive and flat otherwise. Data are from yfinance; metric names follow the "
            "AlphaEdge backtest conventions.",
        )
        note = doc.add_paragraph()
        note.add_run(
            "Disclaimer: For academic use only; not investment advice. Live trading involves liquidity, fees, "
            "regulation, and execution gaps."
        ).italic = True

        doc.add_heading("1. Research Design and Sample Split", level=1)
        _paragraph(
            doc,
            "Training: through 2020-12-31. Validation: calendar year 2021 (grid search for Ridge penalty lambda). "
            "Test: from 2022-01-01 to 2026-04-25. "
            "Alignment: for each Taiwan trading day t, features use the last U.S. trading day strictly before t, "
            "to avoid look-ahead bias.",
        )
        _add_table_from_rows(
            doc,
            ["Dataset Item", "Specification"],
            [
                ["Backtest horizon", "2020-01-01 to 2026-04-25"],
                ["Training set", "2020-01-01 to 2020-12-31"],
                ["Validation set", "2021-01-01 to 2021-12-31"],
                ["Test set", "2022-01-01 to 2026-04-25"],
                ["Data source", "yfinance (daily bars, auto-adjusted)"],
                ["Target and features", "Target: 2330.TW; Features: TSM, ^SOX, TWD=X"],
            ],
            col_widths_cm=[4.5, 12],
        )

        doc.add_heading("2. Strategy Logic", level=1)
        _paragraph(
            doc,
            "(1) Economic rationale: Overnight information is embedded in U.S. prices and FX before the Taiwan "
            "session opens, and may predict the next Taiwan session return for TSMC.",
        )
        _paragraph(
            doc,
            "(2) Timing: For each day d when 2330.TW trades, features are returns from the most recent U.S. session "
            "that ends before d (calendar ordering).",
        )
        _paragraph(
            doc,
            "(3) Variables: Predictors are same-day percentage returns of TSM, SOX, and TWD=X on that U.S. date; "
            "the target is the close-to-close return of 2330.TW on the Taiwan date (daily holding-period proxy).",
        )
        _paragraph(
            doc,
            "(4) Estimation: Ridge regression with intercept; L2 shrinkage on slopes only. Lambda is searched on a log "
            "grid and chosen by validation MSE in 2021; coefficients are re-estimated on training plus validation "
            "before scoring the test period.",
        )
        _paragraph(
            doc,
            "(5) Trading rule: If the predicted return is positive, invest fully in 2330; otherwise hold cash "
            "(zero return, no margin financing modeled).",
        )
        _paragraph(
            doc,
            "(6) Costs: Buy fee when increasing long exposure; sell fee and transaction tax when reducing or exiting, "
            "consistent with the course report assumptions.",
        )
        _paragraph(
            doc,
            "(7) Operational buy/sell timing: when pred_t > 0 and the previous position is flat, the day is treated as "
            "a buy/open-long action; when pred_t <= 0 and the previous position is long, the day is treated as a sell/exit action. "
            "If position state does not change between two consecutive days, no turnover cost is charged.",
        )
        _paragraph(
            doc,
            "(8) Why Ridge is used: to mitigate multicollinearity among correlated predictors (for example TSM and SOX), "
            "shrink unstable coefficients via L2 regularization, and improve out-of-sample robustness while retaining linear interpretability.",
        )

        doc.add_heading("3. Data and Model", level=1)
        _add_table_from_rows(
            doc,
            ["Ticker (yfinance)", "Role"],
            [
                ["TSM", "TSMC ADR, prior U.S. session return"],
                ["2330.TW", "Taiwan listing return; buy-and-hold benchmark"],
                ["^SOX", "Semiconductor sector sentiment / beta proxy"],
                ["TWD=X", "USD/TWD spot (simplified proxy vs. NDF hedging)"],
            ],
            col_widths_cm=[4.5, 12],
        )
        _paragraph(
            doc,
            "The model is Ridge regression. After lambda is fixed, coefficients are fit on training plus validation, "
            "then forecasts and PnL are computed on the held-out test sample.",
        )

        doc.add_heading("4. Methods and Their Purpose", level=1)
        _add_table_from_rows(
            doc,
            ["Method", "Purpose"],
            _method_usage_rows("en"),
            col_widths_cm=[6, 10.5],
        )

        doc.add_heading("5. Test-Set Performance", level=1)
        perf_rows = _format_metrics_rows(meta, metrics_raw, "en")
        if perf_rows:
            _add_table_from_rows(doc, ["Metric", "Value"], perf_rows, col_widths_cm=[5.5, 11])
        else:
            _paragraph(
                doc,
                "metrics_summary.csv not found. Run: .venv/bin/python strategy_lab/run_overnight_signal.py",
            )

        doc.add_heading("6. Charts", level=1)
        _paragraph(
            doc,
            "Figures below are exported as PNG by the pipeline: equity, drawdown, rolling Sharpe, monthly heatmap, "
            "annual IC, rolling IC.",
        )
        figures_en: List[Tuple[str, str]] = [
            ("equity_curve.png", "Equity curve: strategy vs. buy-and-hold 2330.TW"),
            ("mdd_underwater.png", "Drawdown (underwater): strategy vs. benchmark"),
            ("rolling_sharpe.png", "Rolling Sharpe ratio (approx. 63 trading days, annualized)"),
            ("monthly_returns_heatmap.png", "Monthly return heatmap (percent)"),
            ("ic_by_year.png", "Annual information coefficient (IC)"),
            ("rolling_ic.png", "Rolling IC (126 trading-day window)"),
        ]
        for i, (fname, cap) in enumerate(figures_en, start=1):
            _add_figure(doc, _OUTPUT / fname, cap, i, "en")

        doc.add_heading("7. Risks and Limitations", level=1)
        for t in [
            "Linear Ridge may miss nonlinearity and structural breaks; tree models and rolling re-fit are not implemented.",
            "FX uses a spot proxy, not a full NDF hedge.",
            "No limit-up/down, auction microstructure, slippage, or futures margin/roll.",
            "Long-or-flat only; short selling and index futures are out of scope.",
            "Choosing lambda on the validation set implies mild data mining risk.",
        ]:
            doc.add_paragraph(sanitize_for_word(t), style="List Bullet")

        doc.add_heading("8. Reproducibility", level=1)
        _paragraph(
            doc,
            "From the AlphaEdge repo root: run .venv/bin/python strategy_lab/run_overnight_signal.py then "
            ".venv/bin/python strategy_lab/generate_quant_report_docx.py [--lang zh|en|both].",
        )

        doc.add_heading("References", level=1)
        doc.add_paragraph(
            "yfinance; AlphaEdge core/backtest/README.md; course report manuscript (fee and sample assumptions).",
            style="List Bullet",
        )

    doc.save(str(out_doc))
    return out_doc


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate TSMC overnight-signal quant report DOCX.")
    parser.add_argument(
        "--lang",
        choices=("zh", "en", "both"),
        default="both",
        help="Output language (default: both Chinese and English files)",
    )
    args = parser.parse_args()

    paths: List[Path] = []
    if args.lang in ("zh", "both"):
        paths.append(build_report("zh"))
    if args.lang in ("en", "both"):
        paths.append(build_report("en"))

    for p in paths:
        print(f"Generated: {p}")


if __name__ == "__main__":
    main()
